
import re
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_hyperlink(paragraph, text, url):
    # This is a bit complex in python-docx, simplifying to just text for now
    # or a simple run with distinct style if possible.
    # For stability in this environment, I'll just append text.
    # If I wanted real hyperlinks, I'd need to manipulate xml.
    # Let's just output [text](url) style or just the text if it's a TOC link.
    run = paragraph.add_run(text)
    run.font.color.rgb = None # Default
    return run

def apply_markdown_formatting(paragraph, text):
    # Handle bold **text**
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            # Handle links [text](url) - simplified to just text
            # We can try to extract clean text
            link_pattern = r'\[(.*?)\]\(.*?\)'
            clean_text = re.sub(link_pattern, r'\1', part)
            if clean_text != part:
                 paragraph.add_run(clean_text)
            else:
                 paragraph.add_run(part)

def convert_md_to_docx(md_path, docx_path):
    doc = Document()
    
    # Set default font to something standard like Calibri or Arial
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_code_block = False
    
    for line in lines:
        line = line.strip()
        
        if line.startswith('```'):
            in_code_block = not in_code_block
            continue
            
        if in_code_block:
            p = doc.add_paragraph(line)
            p.style = 'No Spacing'
            p.runs[0].font.name = 'Courier New'
            continue
            
        if not line:
            continue
            
        if line.startswith('# '):
            # Title
            p = doc.add_heading(line[2:], 0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '):
            # Heading 1
            doc.add_heading(line[3:], 1)
        elif line.startswith('### '):
            # Heading 2
            doc.add_heading(line[4:], 2)
        elif line.startswith('#### '):
            # Heading 3
            doc.add_heading(line[5:], 3)
        elif line.startswith('- '):
            # List item
            p = doc.add_paragraph(style='List Bullet')
            apply_markdown_formatting(p, line[2:])
        elif line.startswith('1. ') or re.match(r'^\d+\.', line):
            # Numbered list
            p = doc.add_paragraph(style='List Number')
            # Remove the number prefix roughly
            text = re.sub(r'^\d+\.\s+', '', line)
            apply_markdown_formatting(p, text)
        elif line.startswith('|'):
            # Table handling is tricky, skipping or dumping text for now
            # To be safe and simple, dump as monospaced text
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = 'Courier New'
        elif line == '---':
            doc.add_page_break()
        else:
            # Normal paragraph
            p = doc.add_paragraph()
            apply_markdown_formatting(p, line)

    doc.save(docx_path)
    print(f"Successfully created {docx_path}")

if __name__ == "__main__":
    convert_md_to_docx(r"d:\AI Project\Project_Proposal.md", r"d:\AI Project\Project_Proposal.docx")
